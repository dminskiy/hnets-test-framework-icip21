import os, sys
import pandas as pd
import seaborn as sn
import docx
import warnings
import datetime
from docx.shared import Inches
import numpy as np
import math

def find_last(arr, pos_offset):
    pos = len(arr)
    val = arr[pos-1]
    return val, pos+pos_offset

def find_average(x_min, arr, num2average, x_offset):
    sum = 0
    x_min -= x_offset
    if len(arr) < num2average:
        return round(arr[x_min], 2)
    elif len(arr) == num2average:
        for item in arr:
            sum += item
        return round(sum/num2average, 2)
    else:
        if num2average % 2 == 0:
            width = int(num2average/2)
        else:
            width = int((num2average-1)/2)
        add_right = width
        add_left = width

        if x_min - add_left < 0:
            add_right += add_left - x_min
            add_left = x_min
        if x_min + add_right >= len(arr):
            add_right = len(arr) - 1 - x_min
            add_left += len(arr) - 1 - x_min

        pts = 0
        for i in range(-add_left, add_right+1, 1):
            sum += arr[x_min+i]
            pts += 1

        return round(sum/pts, 2)

def find_min_in_array(arr, pos_offset):
    min_val = sys.maxsize
    min_position = -1
    for i in range(len(arr)):
        if arr[i] < min_val:
            min_val = arr[i]
            min_position = i

    return min_val, min_position+pos_offset

def accuracies2errors(accs):
    return [round(100.0 - x, 2) for x in accs]

def get_plot_range(min_number, max_number, samples):

    if len(samples) != (max_number-min_number):
        if (max_number-min_number) % len(samples) == 0:
            step = (max_number-min_number) / len(samples)
        else:
            raise(ValueError("Cannot calculate the range step size for a figure. Number of samples: {}; Min range val:"
                             " {}; Max range val: {}".format(len(samples), min_number, max_number)))
    else:
        step = 1

    step = int(step)

    return range(step, max_number+step, step), step

def plot_errors_pure(plt, report_data):
    epochs = report_data.epochs
    train_errs = report_data.train_accuracies
    test_errs = report_data.test_accuracies
    is_accuracy = True

    if is_accuracy:
        train_errs = accuracies2errors(train_errs)
        test_errs = accuracies2errors(test_errs)

    train_range, train_step = get_plot_range(0, epochs, train_errs)
    test_range, test_step = get_plot_range(0, epochs, test_errs)

    plt.plot(train_range, train_errs, color='blue', linewidth=2, label="Train Error, %")
    plt.plot(test_range, test_errs, color='red', linewidth=2, linestyle=':', label="Test Error, %")

    y_lim = math.floor(max(max(train_errs),max(test_errs))/5)*5+10
    plt.set_ylim(0,y_lim)
    plt.set_yticks(np.arange(0, y_lim, 5))

    return train_errs[len(train_errs)-1]

def plot_errors_with_stats(fig, report_data, plot_top_5=True):
    x_pos_offset = 1

    epochs = report_data.epochs
    test_errs = accuracies2errors(report_data.test_accuracies)

    plt1 = fig.add_subplot(1,1,1)
    #Plot test and train errors
    last_tr_error = plot_errors_pure(plt1, report_data)
    test_range, test_step = get_plot_range(0, epochs, test_errs)

    #Plot additional params - stats
    y_min, x_min = find_min_in_array(test_errs, pos_offset=x_pos_offset)
    min_average_error = find_average(x_min, test_errs, 3, x_offset=x_pos_offset)
    x_min *= test_step

    report_data.min_test_err = y_min
    report_data.min_av_test_err = min_average_error
    report_data.min_test_err_epoch = x_min

    y_last, x_last = find_last(test_errs, pos_offset=x_pos_offset)
    last_average_error = find_average(len(test_errs) - 1, test_errs, 3, x_offset=0)
    x_last *= test_step
    report_data.last_error_av = last_average_error


    plt1.axvline(x=x_min, color='red', label='\nMIN Error:                  {}'
                                             '\nAverage MIN:             {}'
                                             '\nMIN Error at epoch:    {}'.format(y_min, min_average_error, int(x_min)))

    plt1.plot(x_min, y_min, color='red', marker='*', label='\nLast epoch error:     {}'
                                                           '\nAverage Last:          {}'
                                                           '\nLast epoch:              {}'
                                                           '\nLast training error:  {}'.format(y_last, last_average_error, report_data.epochs, last_tr_error))
    #plot top 5 stats
    if report_data.is_available(report_data.top_5_accuracy) and plot_top_5:
        test_top_5_errs = accuracies2errors(report_data.test_accuracies_top_5)
        y5_min, x5_min = find_min_in_array(test_top_5_errs, pos_offset=x_pos_offset)
        min_average_5_error = find_average(x5_min, test_top_5_errs, 3, x_offset=x_pos_offset)
        y5_last, x5_last = find_last(test_top_5_errs, pos_offset=x_pos_offset)
        last_average_5_error = find_average(len(test_top_5_errs) - 1, test_top_5_errs, 3, x_offset=0)
        x5_last *= test_step
        plt1.plot(test_range, test_top_5_errs, color='green', linewidth=2, linestyle=':', label = '\nTest Error, top-5, %')


        plt1.plot(x5_min, y5_min, color='green', marker='*', label='\nMin Top-5 Error:   {}'
                                                                   '\nAverage Min:        {}'
                                                                   '\nFinal top-5 error:  {}'
                                                                   '\nAverage Last:       {}'.format(y5_min,min_average_5_error,y5_last,last_average_5_error))
    plt1.set_title("Plot of Train and Test Errors")
    plt1.legend()

def plot_errors_with_data(report_data, data,fig, title = "", label = ""):
    plt1 = fig.add_subplot(2, 1, 1)
    plot_errors_pure(plt1, report_data)
    plt1.set_title("Test and Train Errors")
    plt1.legend()

    if label == "":
        label = title

    plt2 = fig.add_subplot(2,1,2)
    data_range, _ = get_plot_range(0, len(data), data)
    plt2.plot(data_range, data, color='green', linewidth=2, label=label)
    plt2.set_title(title)
    plt2.legend()

def plot_confusion_matrix(fig, matrix):
    plt = fig.add_subplot(1,1,1)
    df_cm = pd.DataFrame(matrix)
    sn.heatmap(df_cm, annot=True)
    plt.set_title("Confusion Matrix")
    plt.set_xlabel("Target Classes")
    plt.set_ylabel("Predicted Classes")

def save_figure(fig, path_out, name, figures_folder):
    if not (name.endswith(".jpeg") or name.endswith(".jpg") or name.endswith(".png")):
        name += ".jpeg"

    full_path_out = os.path.join(path_out, figures_folder)
    ensure_dir_exists(path_out)
    if path_out != figures_folder:
        ensure_dir_exists(full_path_out)

    full_fig_name = os.path.join(full_path_out, name)
    fig.savefig(full_fig_name)

    print("Figure saved: [{}]".format(full_fig_name))

def generate_report_docx(report_data, doc_path, doc_name, file_in):
    doc = docx.Document()
    doc.add_heading("Network Performance Report", 0)
    st = doc.add_paragraph("Date of the report (YYYY-MM-DD): {}".format(datetime.datetime.now().date().__str__()))
    st.add_run("\nTime of the report: {}".format(datetime.datetime.now().time().__str__().split(".")[0]))
    st.add_run("\nOriginal filename: {}".format(file_in))

    doc.add_heading("Summary", 1)
    summary = doc.add_paragraph("")
    summary.add_run("Dataset: ").bold = True
    summary.add_run("{}".format(report_data.dataset))
    summary.add_run("\nData Used in Training: ").bold = True
    summary.add_run("{}%".format(report_data.data_partition if report_data.data_partition is not None else 100))
    if report_data.is_available(report_data.data_size):
        summary.add_run("Data size: ").bold = True
        summary.add_run("{}x{} pxls".format(report_data.data_size, report_data.data_size))
    summary.add_run("\nDNN: ").bold = True
    summary.add_run("{}".format(report_data.classifier))

    if report_data.scat_enabled and report_data.is_available(report_data.scat_enabled):
        b = summary.add_run("\nScatterNet: ")
        b.bold = True
        summary.add_run("{}".format(report_data.scat_type))
    else:
        b = summary.add_run("\nScatterNet: ")
        b.bold = True
        summary.add_run("Not used")
    summary.add_run("\nEpochs Trained: ").bold = True
    summary.add_run("{}".format(report_data.epochs))

    summary.add_run("\n\nTop-1 Error, %: ").bold = True
    summary.add_run("{}".format(round(100.0 - report_data.final_accuracy,2)))

    if report_data.is_available(report_data.top_5_accuracy):
        summary.add_run("\nTop-5 Error, %: ").bold = True
        summary.add_run("{}".format(round(100.0 - report_data.top_5_accuracy,2)))

    if report_data.is_available(report_data.datetime_start):
        summary.add_run("\n\nTest start date and time: ").bold = True
        summary.add_run("{}; {}".format(report_data.datetime_start.date().__str__(), report_data.datetime_start.time().__str__().split(".")[0]))
    if report_data.is_available(report_data.datetime_end):
        summary.add_run("\nTest end date and time: ").bold = True
        summary.add_run("{}; {}".format(report_data.datetime_end.date().__str__(), report_data.datetime_end.time().__str__().split(".")[0]))

    if report_data.is_available(report_data.total_execution_time):
        summary.add_run("\n\nTotal execution time: ").bold = True
        summary.add_run("{}".format(report_data.total_execution_time))
    if report_data.is_available(report_data.total_tr_time):
        summary.add_run("\nTotal training time:  ").bold = True
        summary.add_run("{}".format(report_data.sec2time(report_data.total_tr_time)))
    if report_data.is_available(report_data.total_tst_time):
        summary.add_run("\nTotal test time:       ").bold = True
        summary.add_run("{}".format(report_data.sec2time(report_data.total_tst_time)))

    if report_data.is_available(report_data.device):
        summary.add_run("\n\nRan on device: ").bold = True
        summary.add_run("{}".format(report_data.device))
        if (report_data.is_available(report_data.num_devices)):
            summary.add_run("\nNumber of devices: ").bold = True
            summary.add_run("{}".format(report_data.num_devices))

    doc.add_heading("Setup", 1)
    doc.add_heading("Data", 2)
    data = doc.add_paragraph("")
    data.add_run("Dataset: ").bold = True
    data.add_run("{}".format(report_data.dataset))
    data.add_run("\nData Used in Training: ").bold = True
    data.add_run("{}%".format(report_data.data_partition if report_data.data_partition is not None else 100))
    if report_data.is_available(report_data.data_size):
        data.add_run("\nData size: ").bold = True
        data.add_run("{}x{} pxl".format(report_data.data_size, report_data.data_size))
    data.add_run("\nAugmentation: ").bold = True
    data.add_run("{}".format(report_data.augment_data))
    data.add_run("\nNumber of classes: ").bold = True
    data.add_run("{}".format(report_data.n_classes))
    data.add_run("\nBatch size: ").bold = True
    data.add_run("{}".format(report_data.batch_sz))

    doc.add_heading("Training Parameters", 2)
    tp = doc.add_paragraph("")
    tp.add_run("Epochs Goal: ").bold = True
    tp.add_run("{}".format(report_data.planned_epochs))
    tp.add_run("\nEpochs Trained: ").bold = True
    tp.add_run("{}".format(report_data.epochs))
    tp.add_run("\nStarting LR: ").bold = True
    tp.add_run("{}".format(report_data.lr))

    if report_data.is_available(report_data.optimiser):
        tp.add_run("\nOptimiser: ").bold = True
        tp.add_run("{}".format(report_data.optimiser))
    else:
        tp.add_run("\nOptimiser: ").bold = True
        tp.add_run("Unknown. Default: SGD")

    tp.add_run("\nCost function: ").bold = True
    if report_data.is_available(report_data.cost_func):
        tp.add_run("{}".format(report_data.cost_func))
    else:
        tp.add_run("Unknown. Default: CrossEntropyLoss")
    tp.add_run("\nLR Scheduler: ").bold = True
    tp.add_run("{}".format(report_data.scheduler_type))
    if report_data.scheduler_type == 'stepLR':
        tp.add_run("\nLR Scheduler step: ").bold = True
        tp.add_run("{}".format(report_data.scheduler_step))
    if report_data.is_available(report_data.net_learnable_params):
        tp.add_run("\nNetwork learnable params: ").bold = True
        tp.add_run("{}".format(report_data.net_learnable_params))
    if report_data.is_available(report_data.scat_learnable_params):
        tp.add_run("\nScattering learnable params: ").bold = True
        tp.add_run("{}".format(report_data.scat_learnable_params))

    doc.add_heading("CNN", 2)
    net = doc.add_paragraph("")
    net.add_run("Network type: ").bold = True
    net.add_run("{}".format(report_data.classifier))
    net.add_run("\nDropout: ").bold = True
    net.add_run("{}".format(report_data.dnn_dropout))

    doc.add_heading("Predefined Filters", 2)
    pf = doc.add_paragraph("")

    if not report_data.scat_enabled:
        pf.add_run("NOT USED").bold = True
    else:
        pf.add_run("Network type: ").bold = True
        pf.add_run("{}".format(report_data.scat_type))
        pf.add_run("\nNetwork order: ").bold = True
        pf.add_run("{}".format(report_data.scat_order))
        pf.add_run("\nJ: ").bold = True
        pf.add_run("{}".format(report_data.J))
        pf.add_run("\nL: ").bold = True
        pf.add_run("{}".format(report_data.L))

    doc.add_heading("Results", 1)
    headings = ["Hybrid", "CNN", "Dataset", "Epochs", "End Error, %", "End Error Av, %", "End Error Top-5, %", "Min Error, %", "Min Error Av, %", "Min Error Epoch"]
    hybrid_part = report_data.scat_type if report_data.scat_enabled else "None"
    top_5 = round(100.0-report_data.top_5_accuracy,2) if report_data.is_available(report_data.top_5_accuracy) else "None"
    table_data = [hybrid_part, report_data.classifier, report_data.dataset, report_data.epochs, round(100.-report_data.final_accuracy, 2), report_data.last_error_av,
                top_5, report_data.min_test_err, report_data.min_av_test_err, report_data.min_test_err_epoch]
    table = doc.add_table(rows = 2, cols = len(headings))
    table.style = 'TableGrid'
    head_row = table.rows[0].cells
    data_row = table.rows[1].cells
    for i in range(len(head_row)):
        head_row[i].text = headings[i]
        data_row[i].text = str(table_data[i])

    doc.add_heading("Graphs", 1)

    pic_width_inch = Inches(7.5)
    if report_data.is_available(report_data.img_path_detailed_errs):
        doc.add_heading("Train and Test Errors",2)
        doc.add_picture(report_data.img_path_detailed_errs, width=pic_width_inch)
    if report_data.is_available(report_data.img_path_losses_n_errs):
        doc.add_heading("Loss function output",2)
        doc.add_picture(report_data.img_path_losses_n_errs, width=pic_width_inch)
    if report_data.is_available(report_data.img_path_lrs_n_errs):
        doc.add_heading("Changes in LR",2)
        doc.add_picture(report_data.img_path_lrs_n_errs, width=pic_width_inch)
    if report_data.is_available(report_data.img_path_confusion_mat):
        doc.add_heading("Confusion Matrix",2)
        doc.add_picture(report_data.img_path_confusion_mat, width=pic_width_inch)

    doc.add_heading("Conclusions", 1)
    b = doc.add_paragraph("TO BE ADDED MANUALLY")
    b.bold = True

    ensure_dir_exists(doc_path)

    try:
        doc.save(os.path.join(doc_path, doc_name))
        print("The performance report was saved @ [{}]".format(os.path.join(doc_path, doc_name)))
    except:
        warnings.warn("Error while saving the report document: [{}]".format(os.path.join(doc_path, doc_name)), Warning)